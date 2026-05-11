from docx import Document
from docx.shared import Pt, RGBColor

from rewriter.processor import process_document


def test_rewriteable_paragraph_text_is_replaced_and_formatting_preserved(
    sample_docx_single_run_para, tmp_path
):
    output = tmp_path / "out.docx"
    fake_rewrite = "改写后的段落文字，长度适当且保留所有原始信息。"

    report = process_document(
        sample_docx_single_run_para,
        output,
        qwen_call=lambda system, user: fake_rewrite,
    )

    assert report.rewritten == 1
    assert report.api_failures == []

    out_doc = Document(str(output))
    para = out_doc.paragraphs[0]
    assert para.text == fake_rewrite
    run = para.runs[0]
    assert run.font.name == "宋体"
    assert run.font.size == Pt(12)
    assert run.font.color.rgb == RGBColor(0x00, 0x00, 0x00)


def test_skipped_paragraphs_keep_original_text(sample_docx, tmp_path):
    output = tmp_path / "out.docx"
    sentinel = "[REWRITTEN_BY_TEST]"

    process_document(
        sample_docx,
        output,
        qwen_call=lambda system, user: sentinel * 5,
    )

    out_doc = Document(str(output))
    out_texts = [p.text for p in out_doc.paragraphs]

    titles_unchanged = [
        "深度学习在图像识别中的应用研究",
        "第一章 绪论",
        "图1 系统总体架构示意图",
        "表1 实验参数设置",
        "参考文献",
    ]
    for title in titles_unchanged:
        assert title in out_texts, f"{title!r} should be preserved"

    assert "如下所示。" in out_texts
    assert any(t.startswith("[1] 张三") for t in out_texts)
    assert any("pure English" in t for t in out_texts)


def test_report_records_skip_reasons_and_rewrite_count(sample_docx, tmp_path):
    output = tmp_path / "out.docx"

    def fake_rewrite(system, user):
        original = user.split("\n\n", 1)[1]
        return original.replace("。", "！", 1)

    report = process_document(sample_docx, output, qwen_call=fake_rewrite)

    assert report.total_paragraphs == 12
    assert report.rewritten == 2
    assert "heading" in report.skipped_by_reason
    assert "caption_prefix" in report.skipped_by_reason
    assert "reference_entry" in report.skipped_by_reason or "after_references" in report.skipped_by_reason
    assert "mixed_format" in report.skipped_by_reason
    assert "too_short" in report.skipped_by_reason
    assert "no_chinese" in report.skipped_by_reason


def test_progress_callback_invoked_with_completion_counts(
    sample_docx_single_run_para, tmp_path
):
    output = tmp_path / "out.docx"
    events = []

    process_document(
        sample_docx_single_run_para,
        output,
        qwen_call=lambda s, u: "改写后的段落文字，长度适当。",
        progress=lambda done, total: events.append((done, total)),
    )

    assert events == [(1, 1)]


def test_progress_callback_handles_multiple_paragraphs(tmp_path):
    doc = Document()
    for i in range(3):
        doc.add_paragraph(f"这是第{i}段需要改写的正文，长度足够触发改写规则。")
    src = tmp_path / "three.docx"
    doc.save(str(src))

    output = tmp_path / "out.docx"
    events = []

    process_document(
        src,
        output,
        qwen_call=lambda s, u: "改写后段落，长度合适并不漂移。",
        max_workers=2,
        progress=lambda done, total: events.append((done, total)),
    )

    assert len(events) == 3
    assert all(total == 3 for _, total in events)
    assert [done for done, _ in events] == [1, 2, 3]
