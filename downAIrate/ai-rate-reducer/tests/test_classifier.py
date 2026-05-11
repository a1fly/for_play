from docx import Document
from rewriter.classifier import Classifier, Classification


def test_heading_1_is_skipped(sample_docx):
    doc = Document(str(sample_docx))
    classifier = Classifier()
    # Paragraph 0 in the fixture is a Heading 1
    result = classifier.classify(doc.paragraphs[0])
    assert result == Classification(rewrite=False, skip_reason="heading")


def test_figure_caption_is_skipped(sample_docx):
    doc = Document(str(sample_docx))
    classifier = Classifier()
    for p in doc.paragraphs:
        if p.text.startswith("图1"):
            result = classifier.classify(p)
            assert result == Classification(rewrite=False, skip_reason="caption_prefix")
            return
    raise AssertionError("figure caption not found in fixture")


def test_table_caption_is_skipped(sample_docx):
    doc = Document(str(sample_docx))
    classifier = Classifier()
    for p in doc.paragraphs:
        if p.text.startswith("表1"):
            result = classifier.classify(p)
            assert result == Classification(rewrite=False, skip_reason="caption_prefix")
            return
    raise AssertionError("table caption not found in fixture")


def test_reference_entry_starting_with_bracket_is_skipped(sample_docx):
    doc = Document(str(sample_docx))
    classifier = Classifier()
    # Need to feed all paragraphs in order so the references flag triggers
    results = [classifier.classify(p) for p in doc.paragraphs]
    # Find the [1] paragraph
    for p, result in zip(doc.paragraphs, results):
        if p.text.startswith("[1]"):
            assert result.rewrite is False
            assert result.skip_reason in ("reference_entry", "after_references")
            return
    raise AssertionError("reference entry not found")


def test_paragraph_after_references_heading_is_skipped(sample_docx):
    doc = Document(str(sample_docx))
    classifier = Classifier()
    results = [classifier.classify(p) for p in doc.paragraphs]
    # Last paragraph in fixture is normal-looking body text AFTER references heading
    last_text_para = doc.paragraphs[-1]
    last_result = results[-1]
    assert "参考文献区之后" in last_text_para.text
    assert last_result == Classification(rewrite=False, skip_reason="after_references")


def test_pure_english_paragraph_is_skipped(sample_docx):
    doc = Document(str(sample_docx))
    classifier = Classifier()
    for p in doc.paragraphs:
        if "pure English" in p.text:
            result = classifier.classify(p)
            assert result.rewrite is False
            assert result.skip_reason in ("no_chinese", "low_chinese_ratio")
            return
    raise AssertionError("pure English paragraph not found")
