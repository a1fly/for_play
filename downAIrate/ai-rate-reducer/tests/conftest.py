import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor


@pytest.fixture
def sample_docx(tmp_path):
    """Build a fixture docx covering every classifier rule.

    Returns the path. Paragraph indexes are documented inline so tests
    can refer to them by index.
    """
    doc = Document()

    # 0: Title (skip — Heading 1)
    doc.add_heading("深度学习在图像识别中的应用研究", level=1)

    # 1: Heading 2 (skip)
    doc.add_heading("第一章 绪论", level=2)

    # 2: Normal body paragraph, single run, Chinese >= 20%, len >= 15 (rewrite)
    doc.add_paragraph(
        "近年来人工智能技术得到了广泛应用，本文将围绕深度学习方法在图像识别任务中的实现展开讨论。"
    )

    # 3: Short paragraph (skip — too short, < 15 chars)
    doc.add_paragraph("如下所示。")

    # 4: Pure English (skip — no Chinese)
    doc.add_paragraph("This is a pure English caption that should be skipped.")

    # 5: Mixed format paragraph (skip — bold word inside)
    p_mixed = doc.add_paragraph("研究表明")
    run_bold = p_mixed.add_run("深度学习")
    run_bold.bold = True
    p_mixed.add_run("在多个领域都有突出表现，本段是混合格式段落。")

    # 6: Figure caption (skip — starts with "图1")
    doc.add_paragraph("图1 系统总体架构示意图")

    # 7: Table caption (skip — starts with "表1")
    doc.add_paragraph("表1 实验参数设置")

    # 8: Another rewriteable body paragraph
    doc.add_paragraph(
        "卷积神经网络通过逐层提取图像特征，能够有效完成分类任务，这是当前主流方法。"
    )

    # 9: References heading (skip — heading) — triggers references-zone flag
    doc.add_heading("参考文献", level=1)

    # 10: Reference entry, starts with [1] (skip)
    doc.add_paragraph("[1] 张三. 深度学习概论[M]. 北京: 科学出版社, 2020.")

    # 11: After references, even a long Chinese paragraph should be skipped
    doc.add_paragraph(
        "这是参考文献区之后的一段长文字，按规则应被跳过即使内容像正文。"
    )

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_docx_single_run_para(tmp_path):
    """A docx with one rewriteable paragraph whose run has specific formatting
    we can assert is preserved after rewrite."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(
        "这是一段需要被改写的正文文字，长度足够触发改写规则。"
    )
    run.font.name = "宋体"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    path = tmp_path / "single_run.docx"
    doc.save(str(path))
    return path
